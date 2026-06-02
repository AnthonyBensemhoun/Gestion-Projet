"""
Atelier - Gestion de projet collaborative
Backend FastAPI + SQLite
"""
from fastapi import FastAPI, Depends, HTTPException, Request, Form, UploadFile, File, Cookie, Response
from fastapi.responses import HTMLResponse, RedirectResponse, FileResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from sqlmodel import SQLModel, Field, Session, create_engine, select
from sqlalchemy import text, update as sa_update, LargeBinary, Column
from urllib.parse import quote as _urlquote
from typing import Optional, List
from datetime import date, datetime, timedelta
from passlib.hash import bcrypt
import secrets, os, json, io, re, smtplib, threading
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from apscheduler.schedulers.background import BackgroundScheduler
try:
    import anthropic
except ImportError:
    anthropic = None

# ---------------- Config ----------------
DB_URL = os.environ.get("DATABASE_URL", "sqlite:///./atelier.db")
SECRET_ADMIN_TOKEN = os.environ.get("ADMIN_BOOTSTRAP_TOKEN", "atelier-setup")
if DB_URL.startswith("postgresql://") or DB_URL.startswith("postgres://"):
    import ssl as _ssl
    DB_URL = DB_URL.replace("postgresql://", "postgresql+pg8000://", 1).replace("postgres://", "postgresql+pg8000://", 1)
    _ssl_ctx = _ssl.create_default_context()
    _ssl_ctx.check_hostname = False
    _ssl_ctx.verify_mode = _ssl.CERT_NONE
    engine = create_engine(DB_URL, connect_args={"ssl_context": _ssl_ctx}, pool_pre_ping=True)
else:
    engine = create_engine(DB_URL, connect_args={"check_same_thread": False})

# ---------------- Modèles ----------------
class User(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    email: str = Field(unique=True, index=True)
    name: str
    password_hash: str
    role: str = "user"   # 'admin' ou 'user'
    created_at: datetime = Field(default_factory=datetime.utcnow)
    must_change_password: bool = Field(default=False)
    last_login: Optional[datetime] = Field(default=None)
    last_seen: Optional[datetime] = Field(default=None)

class Session_(SQLModel, table=True):
    token: str = Field(primary_key=True)
    user_id: int
    expires: datetime

class Project(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    name: str
    description: str = ""
    created_at: datetime = Field(default_factory=datetime.utcnow)

class Task(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    project_id: int = Field(index=True)
    title: str
    description: str = ""
    assignee_id: Optional[int] = None
    priority: str = "m"
    start_date: Optional[date] = None
    due_date: Optional[date] = None
    status: str = "todo"
    progress: int = 0
    estimated_hours: Optional[float] = None
    actual_hours: Optional[float] = None
    milestone_id: Optional[int] = None

class Absence(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    user_id: int
    kind: str
    from_date: date
    to_date: date

class AckedAlert(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    alert_key: str = Field(index=True)  # ex: "late:42" — task_id concerné

class SubTask(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    task_id: int = Field(index=True)
    title: str
    done: bool = Field(default=False)
    position: int = Field(default=0)

class Comment(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    task_id: int = Field(index=True)
    user_id: int
    text: str
    created_at: datetime = Field(default_factory=datetime.utcnow)

class Tag(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    project_id: int = Field(index=True)
    name: str
    color: str = "#e8642f"

class TaskTag(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    task_id: int = Field(index=True)
    tag_id: int = Field(index=True)

class Milestone(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    project_id: int = Field(index=True)
    name: str
    due_date: Optional[date] = None
    description: str = ""

class Setting(SQLModel, table=True):
    key: str = Field(primary_key=True)
    value: str

class Document(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    project_id: Optional[int] = Field(default=None, index=True)  # None = document de service (global)
    name: str
    description: str = ""
    doc_type: str = "DOC"      # SOP / PROTO / REPORT / FORM / IT / DOC
    reference: str = ""        # référence auto : ex SOP-2026-014
    status: str = "draft"      # draft / review / approved (dérivé de la phase)
    phase: str = "redaction"   # workflow : redaction / revue_equipe / revue_qa / approbation / pret_qms
    phase_since: Optional[datetime] = Field(default=None)  # entrée dans la phase actuelle (SLA)
    assigned_to: Optional[int] = Field(default=None)  # personne chez qui le document se trouve actuellement
    locked_by: Optional[int] = Field(default=None)   # user_id ayant verrouillé pour édition
    locked_at: Optional[datetime] = Field(default=None)
    created_by: Optional[int] = Field(default=None)
    created_at: datetime = Field(default_factory=datetime.utcnow)

class DocSignature(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    document_id: int = Field(index=True)
    user_id: Optional[int] = Field(default=None)
    user_name: str = ""        # snapshot immuable du signataire
    phase: str = ""            # phase signée
    meaning: str = ""          # signification (ex : Approbation, Prêt pour QMS)
    reason: str = ""           # motif saisi
    version: int = 0           # n° de version signé
    signed_at: datetime = Field(default_factory=datetime.utcnow)

class DocAck(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    document_id: int = Field(index=True)
    user_id: int = Field(index=True)
    version: int = 0           # version accusée
    acknowledged_at: datetime = Field(default_factory=datetime.utcnow)

class DocComment(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    document_id: int = Field(index=True)
    user_id: Optional[int] = Field(default=None)
    text: str = ""
    created_at: datetime = Field(default_factory=datetime.utcnow)

class DocWorkflowEvent(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    document_id: int = Field(index=True)
    phase: str = ""
    assigned_to: Optional[int] = Field(default=None)
    moved_by: Optional[int] = Field(default=None)
    note: str = ""
    created_at: datetime = Field(default_factory=datetime.utcnow)

class DocumentVersion(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    document_id: int = Field(index=True)
    version: int = 1
    filename: str = ""
    mime_type: str = ""
    size: int = 0
    content: bytes = Field(sa_column=Column(LargeBinary))
    uploaded_by: Optional[int] = Field(default=None)
    uploaded_at: datetime = Field(default_factory=datetime.utcnow)
    note: str = ""

class AuditLog(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    user_name: str = ""
    action: str = ""
    details: str = ""
    created_at: datetime = Field(default_factory=datetime.utcnow)

class Notification(SQLModel, table=True):
    id: Optional[int] = Field(default=None, primary_key=True)
    user_id: int = Field(index=True)        # destinataire
    kind: str = ""                          # doc_assigned / mention / ...
    title: str = ""
    body: str = ""
    doc_id: Optional[int] = Field(default=None)
    task_id: Optional[int] = Field(default=None)
    read: bool = Field(default=False)
    created_at: datetime = Field(default_factory=datetime.utcnow)

# ---------------- Helpers ----------------
def get_session():
    with Session(engine) as s:
        yield s

def _setting(s: Session, key: str, default: str = "") -> str:
    r = s.get(Setting, key)
    return r.value if r else default

def _set_setting(s: Session, key: str, value: str):
    r = s.get(Setting, key)
    if r:
        r.value = value
    else:
        r = Setting(key=key, value=value)
    s.add(r); s.commit()

def get_current_user(request: Request, s: Session = Depends(get_session)) -> Optional[User]:
    token = request.cookies.get("atelier_session")
    if not token:
        return None
    sess = s.get(Session_, token)
    if not sess or sess.expires < datetime.utcnow():
        return None
    return s.get(User, sess.user_id)

def _audit(s: Session, user_name: str, action: str, details: str = ""):
    s.add(AuditLog(user_name=user_name, action=action, details=details))

def _notify(s: Session, user_id, kind, title, body="", doc_id=None, task_id=None, actor_id=None):
    """Crée une notification in-app (sauf si destinataire == acteur)."""
    if not user_id or user_id == actor_id:
        return
    s.add(Notification(user_id=user_id, kind=kind, title=title, body=body,
                       doc_id=doc_id, task_id=task_id))

def _update_last_seen(user_id: int):
    try:
        with Session(engine) as s:
            s.execute(sa_update(User).where(User.id == user_id).values(last_seen=datetime.utcnow()))
            s.commit()
    except Exception:
        pass

def require_user(request: Request, s: Session = Depends(get_session)) -> User:
    u = get_current_user(request, s)
    if not u:
        raise HTTPException(401, "Non connecté")
    now = datetime.utcnow()
    if not u.last_seen or (now - u.last_seen).total_seconds() > 300:
        threading.Thread(target=_update_last_seen, args=(u.id,), daemon=True).start()
    return u

def require_admin(request: Request, s: Session = Depends(get_session)) -> User:
    u = require_user(request, s)
    if u.role != "admin":
        raise HTTPException(403, "Accès réservé aux administrateurs")
    return u

# ---------------- App ----------------
app = FastAPI(title="Atelier")
templates = Jinja2Templates(directory="templates")
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.on_event("startup")
def on_startup():
    SQLModel.metadata.create_all(engine)
    with engine.connect() as conn:
        migrations = [
            ('ALTER TABLE "user" ADD COLUMN IF NOT EXISTS must_change_password BOOLEAN DEFAULT FALSE NOT NULL',
             "ALTER TABLE user ADD COLUMN must_change_password BOOLEAN DEFAULT 0 NOT NULL"),
            ('ALTER TABLE task ADD COLUMN IF NOT EXISTS estimated_hours FLOAT',
             "ALTER TABLE task ADD COLUMN estimated_hours FLOAT"),
            ('ALTER TABLE task ADD COLUMN IF NOT EXISTS actual_hours FLOAT',
             "ALTER TABLE task ADD COLUMN actual_hours FLOAT"),
            ('ALTER TABLE task ADD COLUMN IF NOT EXISTS milestone_id INTEGER',
             "ALTER TABLE task ADD COLUMN milestone_id INTEGER"),
            ('ALTER TABLE "user" ADD COLUMN IF NOT EXISTS last_login TIMESTAMP',
             "ALTER TABLE user ADD COLUMN last_login TIMESTAMP"),
            ('ALTER TABLE "user" ADD COLUMN IF NOT EXISTS last_seen TIMESTAMP',
             "ALTER TABLE user ADD COLUMN last_seen TIMESTAMP"),
            ("ALTER TABLE document ADD COLUMN IF NOT EXISTS phase VARCHAR DEFAULT 'redaction'",
             "ALTER TABLE document ADD COLUMN phase VARCHAR DEFAULT 'redaction'"),
            ('ALTER TABLE document ADD COLUMN IF NOT EXISTS assigned_to INTEGER',
             "ALTER TABLE document ADD COLUMN assigned_to INTEGER"),
            ("ALTER TABLE document ADD COLUMN IF NOT EXISTS doc_type VARCHAR DEFAULT 'DOC'",
             "ALTER TABLE document ADD COLUMN doc_type VARCHAR DEFAULT 'DOC'"),
            ("ALTER TABLE document ADD COLUMN IF NOT EXISTS reference VARCHAR DEFAULT ''",
             "ALTER TABLE document ADD COLUMN reference VARCHAR DEFAULT ''"),
            ('ALTER TABLE document ADD COLUMN IF NOT EXISTS phase_since TIMESTAMP',
             "ALTER TABLE document ADD COLUMN phase_since TIMESTAMP"),
        ]
        for pg_stmt, sq_stmt in migrations:
            try:
                stmt = sq_stmt if DB_URL.startswith("sqlite") else pg_stmt
                conn.execute(text(stmt)); conn.commit()
            except Exception:
                pass
    # créer un admin par défaut si aucun n'existe
    with Session(engine) as s:
        any_admin = s.exec(select(User).where(User.role == "admin")).first()
        if not any_admin:
            admin_email = os.environ.get("ADMIN_EMAIL", "admin@atelier.local")
            admin_password = os.environ.get("ADMIN_PASSWORD", "admin")
            u = User(email=admin_email, name="Administrateur",
                     password_hash=bcrypt.hash(admin_password), role="admin")
            s.add(u); s.commit()
            print(f"[Atelier] Admin créé : {admin_email} / {admin_password}")
        if not _setting(s, "app_name"):
            _set_setting(s, "app_name", "Helix")
        # Rebranding unique vers "Helix" (ne s'applique qu'une seule fois)
        if not _setting(s, "branding_helix"):
            _set_setting(s, "app_name", "Helix")
            _set_setting(s, "branding_helix", "1")

# ---------------- Pages ----------------
@app.get("/", response_class=HTMLResponse)
def home(request: Request, s: Session = Depends(get_session)):
    u = get_current_user(request, s)
    if not u:
        return RedirectResponse("/login")
    return templates.TemplateResponse("app.html", {
        "request": request, "user": u,
        "app_name": _setting(s, "app_name", "Helix"),
        "app_logo": _setting(s, "app_logo", ""),
        "is_admin": u.role == "admin",
        "must_change_password": u.must_change_password
    })

@app.get("/login", response_class=HTMLResponse)
def login_page(request: Request, s: Session = Depends(get_session)):
    return templates.TemplateResponse("login.html", {
        "request": request,
        "app_name": _setting(s, "app_name", "Helix"),
        "app_logo": _setting(s, "app_logo", "")
    })

@app.post("/login")
def login(response: Response, email: str = Form(...), password: str = Form(...),
          s: Session = Depends(get_session)):
    user = s.exec(select(User).where(User.email == email.lower().strip())).first()
    if not user or not bcrypt.verify(password, user.password_hash):
        return RedirectResponse("/login?error=1", status_code=303)
    now = datetime.utcnow()
    token = secrets.token_urlsafe(32)
    s.add(Session_(token=token, user_id=user.id, expires=now + timedelta(days=30)))
    s.execute(sa_update(User).where(User.id == user.id).values(last_login=now, last_seen=now))
    _audit(s, user.name, "Connexion", f"Email : {user.email}")
    s.commit()
    resp = RedirectResponse("/", status_code=303)
    resp.set_cookie("atelier_session", token, httponly=True, samesite="lax", max_age=30*86400)
    return resp

@app.get("/logout")
def logout(request: Request, s: Session = Depends(get_session)):
    token = request.cookies.get("atelier_session")
    if token:
        sess = s.get(Session_, token)
        if sess:
            s.delete(sess); s.commit()
    resp = RedirectResponse("/login", status_code=303)
    resp.delete_cookie("atelier_session")
    return resp

# ---------------- API : Users / Team ----------------
@app.get("/api/me")
def api_me(u: User = Depends(require_user)):
    return {"id": u.id, "name": u.name, "email": u.email, "role": u.role, "must_change_password": u.must_change_password}

@app.get("/api/users")
def list_users(u: User = Depends(require_user), s: Session = Depends(get_session)):
    now = datetime.utcnow()
    return [{"id": x.id, "name": x.name, "email": x.email, "role": x.role,
             "last_login": x.last_login.isoformat() if x.last_login else None,
             "online": bool(x.last_seen and (now - x.last_seen).total_seconds() < 900)}
            for x in s.exec(select(User).order_by(User.name)).all()]

@app.post("/api/users")
def create_user(data: dict, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    email = data.get("email", "").lower().strip()
    name = data.get("name", "").strip()
    password = data.get("password") or "123456"
    if not email or not name:
        raise HTTPException(400, "Nom et email requis")
    if s.exec(select(User).where(User.email == email)).first():
        raise HTTPException(400, "Cet email existe déjà")
    new = User(email=email, name=name, password_hash=bcrypt.hash(password),
               role=data.get("role", "user"), must_change_password=True)
    s.add(new)
    _audit(s, u.name, "Création utilisateur", f"{name} ({email})")
    s.commit(); s.refresh(new)
    return {"id": new.id, "name": new.name, "email": new.email, "role": new.role,
            "initial_password": password}

@app.put("/api/users/{uid}")
def update_user(uid: int, data: dict, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    target = s.get(User, uid)
    if not target:
        raise HTTPException(404, "Utilisateur introuvable")
    updates = {}
    if "name" in data and data["name"].strip():
        updates["name"] = data["name"].strip()
    if "email" in data:
        new_email = data["email"].lower().strip()
        if new_email != target.email:
            conflict = s.exec(select(User).where(User.email == new_email)).first()
            if conflict:
                raise HTTPException(400, "Cet email est déjà utilisé par un autre compte")
        updates["email"] = new_email
    if "role" in data and data["role"] in ("admin", "user"):
        updates["role"] = data["role"]
    if data.get("password"):
        updates["password_hash"] = bcrypt.hash(data["password"])
    if updates:
        s.execute(sa_update(User).where(User.id == uid).values(**updates))
        _audit(s, u.name, "Modification utilisateur", f"Utilisateur ID {uid} — champs : {list(updates.keys())}")
        s.commit()
    return {"ok": True}

@app.delete("/api/users/{uid}")
def delete_user(uid: int, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    if uid == u.id:
        raise HTTPException(400, "Tu ne peux pas te supprimer toi-même")
    target = s.get(User, uid)
    if target:
        _audit(s, u.name, "Suppression utilisateur", f"{target.name} ({target.email})")
        for t in s.exec(select(Task).where(Task.assignee_id == uid)).all():
            t.assignee_id = None; s.add(t)
        for a in s.exec(select(Absence).where(Absence.user_id == uid)).all():
            s.delete(a)
        for sess in s.exec(select(Session_).where(Session_.user_id == uid)).all():
            s.delete(sess)
        s.delete(target); s.commit()
    return {"ok": True}

@app.put("/api/me/password")
def change_my_password(data: dict, u: User = Depends(require_user), s: Session = Depends(get_session)):
    new_pw = data.get("password", "").strip()
    if len(new_pw) < 6:
        raise HTTPException(400, "Le mot de passe doit faire au moins 6 caractères")
    target = s.get(User, u.id)
    target.password_hash = bcrypt.hash(new_pw)
    target.must_change_password = False
    s.add(target); s.commit()
    return {"ok": True}

# ---------------- API : Projects ----------------
@app.get("/api/projects")
def list_projects(u: User = Depends(require_user), s: Session = Depends(get_session)):
    return [{"id": p.id, "name": p.name, "description": p.description}
            for p in s.exec(select(Project).order_by(Project.name)).all()]

@app.post("/api/projects")
def create_project(data: dict, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    name = data.get("name", "").strip()
    if not name: raise HTTPException(400, "Nom requis")
    p = Project(name=name, description=data.get("description", "").strip())
    s.add(p)
    _audit(s, u.name, "Création projet", name)
    s.commit(); s.refresh(p)
    return {"id": p.id, "name": p.name, "description": p.description}

@app.put("/api/projects/{pid}")
def update_project(pid: int, data: dict, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    p = s.get(Project, pid)
    if not p: raise HTTPException(404)
    if "name" in data: p.name = data["name"].strip()
    if "description" in data: p.description = data["description"].strip()
    _audit(s, u.name, "Modification projet", p.name)
    s.add(p); s.commit()
    return {"ok": True}

@app.delete("/api/projects/{pid}")
def delete_project(pid: int, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    p = s.get(Project, pid)
    if p:
        _audit(s, u.name, "Suppression projet", p.name)
        for t in s.exec(select(Task).where(Task.project_id == pid)).all():
            s.delete(t)
        s.delete(p); s.commit()
    return {"ok": True}

# ---------------- API : Tasks ----------------
def task_dict(t: Task) -> dict:
    return {
        "id": t.id, "project_id": t.project_id, "title": t.title,
        "description": t.description, "assignee_id": t.assignee_id,
        "priority": t.priority,
        "start_date": t.start_date.isoformat() if t.start_date else None,
        "due_date": t.due_date.isoformat() if t.due_date else None,
        "estimated_hours": t.estimated_hours, "actual_hours": t.actual_hours,
        "milestone_id": t.milestone_id,
        "status": t.status, "progress": t.progress
    }

@app.get("/api/tasks")
def list_tasks(project_id: Optional[int] = None,
               u: User = Depends(require_user), s: Session = Depends(get_session)):
    q = select(Task)
    if project_id is not None:
        q = q.where(Task.project_id == project_id)
    return [task_dict(t) for t in s.exec(q).all()]

def _parse_date(v):
    if not v: return None
    return date.fromisoformat(v) if isinstance(v, str) else v

@app.post("/api/tasks")
def create_task(data: dict, u: User = Depends(require_user), s: Session = Depends(get_session)):
    # Tout utilisateur peut créer une tâche (assignée à lui-même par défaut)
    title = data.get("title", "").strip()
    if not title: raise HTTPException(400, "Titre requis")
    pid = data.get("project_id")
    if not pid or not s.get(Project, pid):
        raise HTTPException(400, "Projet invalide")
    assignee = data.get("assignee_id")
    # Si non admin : assignation forcée à soi-même
    if u.role != "admin":
        assignee = u.id
    t = Task(
        project_id=pid, title=title,
        description=data.get("description", ""),
        assignee_id=assignee, priority=data.get("priority", "m"),
        start_date=_parse_date(data.get("start_date")),
        due_date=_parse_date(data.get("due_date")),
        status=data.get("status", "todo"),
        progress=int(data.get("progress", 0))
    )
    _audit(s, u.name, "Création tâche", f"'{title}' — projet {pid}")
    s.add(t); s.commit(); s.refresh(t)
    return task_dict(t)

def _can_edit_task(u: User, t: Task) -> bool:
    return u.role == "admin" or t.assignee_id == u.id

@app.put("/api/tasks/{tid}")
def update_task(tid: int, data: dict, u: User = Depends(require_user), s: Session = Depends(get_session)):
    t = s.get(Task, tid)
    if not t: raise HTTPException(404)
    if not _can_edit_task(u, t):
        raise HTTPException(403, "Tu ne peux modifier que tes propres tâches")
    # Les non-admins ne peuvent pas changer l'assignation ni le projet
    if u.role != "admin":
        data.pop("assignee_id", None)
        data.pop("project_id", None)
    for k in ("title", "description", "priority", "status"):
        if k in data: setattr(t, k, data[k])
    if "progress" in data: t.progress = int(data["progress"])
    if "start_date" in data: t.start_date = _parse_date(data["start_date"])
    if "due_date" in data: t.due_date = _parse_date(data["due_date"])
    if "assignee_id" in data: t.assignee_id = data["assignee_id"]
    if "project_id" in data and s.get(Project, data["project_id"]): t.project_id = data["project_id"]
    if "estimated_hours" in data: t.estimated_hours = float(data["estimated_hours"]) if data["estimated_hours"] else None
    if "actual_hours" in data: t.actual_hours = float(data["actual_hours"]) if data["actual_hours"] else None
    if "milestone_id" in data: t.milestone_id = data["milestone_id"]
    if t.status == "done": t.progress = 100
    _audit(s, u.name, "Modification tâche", f"'{t.title}' — statut : {t.status}")
    s.add(t); s.commit()
    # toute mise à jour invalide les acquittements liés à cette tâche
    for ack in s.exec(select(AckedAlert).where(AckedAlert.alert_key.endswith(f":{tid}"))).all():
        s.delete(ack)
    s.commit()
    return task_dict(t)

@app.delete("/api/tasks/{tid}")
def delete_task(tid: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    t = s.get(Task, tid)
    if t:
        if not _can_edit_task(u, t):
            raise HTTPException(403, "Tu ne peux supprimer que tes propres tâches")
        _audit(s, u.name, "Suppression tâche", f"'{t.title}'")
        s.delete(t); s.commit()
    return {"ok": True}

# ---------------- API : Sous-tâches ----------------
@app.get("/api/tasks/{tid}/subtasks")
def list_subtasks(tid: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    return [{"id": st.id, "title": st.title, "done": st.done, "position": st.position}
            for st in s.exec(select(SubTask).where(SubTask.task_id == tid).order_by(SubTask.position)).all()]

@app.post("/api/tasks/{tid}/subtasks")
def create_subtask(tid: int, data: dict, u: User = Depends(require_user), s: Session = Depends(get_session)):
    t = s.get(Task, tid)
    if not t: raise HTTPException(404)
    if not _can_edit_task(u, t): raise HTTPException(403)
    title = data.get("title", "").strip()
    if not title: raise HTTPException(400, "Titre requis")
    pos = s.exec(select(SubTask).where(SubTask.task_id == tid)).all()
    st = SubTask(task_id=tid, title=title, position=len(pos))
    s.add(st); s.commit(); s.refresh(st)
    return {"id": st.id, "title": st.title, "done": st.done, "position": st.position}

@app.put("/api/subtasks/{sid}")
def update_subtask(sid: int, data: dict, u: User = Depends(require_user), s: Session = Depends(get_session)):
    st = s.get(SubTask, sid)
    if not st: raise HTTPException(404)
    t = s.get(Task, st.task_id)
    if not _can_edit_task(u, t): raise HTTPException(403)
    if "title" in data: st.title = data["title"].strip()
    if "done" in data: st.done = bool(data["done"])
    s.add(st); s.commit()
    return {"ok": True}

@app.delete("/api/subtasks/{sid}")
def delete_subtask(sid: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    st = s.get(SubTask, sid)
    if st:
        t = s.get(Task, st.task_id)
        if not _can_edit_task(u, t): raise HTTPException(403)
        s.delete(st); s.commit()
    return {"ok": True}

# ---------------- API : Commentaires ----------------
@app.get("/api/tasks/{tid}/comments")
def list_comments(tid: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    users = {x.id: x.name for x in s.exec(select(User)).all()}
    return [{"id": c.id, "user_id": c.user_id, "author": users.get(c.user_id, "?"),
             "text": c.text, "created_at": c.created_at.isoformat()}
            for c in s.exec(select(Comment).where(Comment.task_id == tid).order_by(Comment.created_at)).all()]

@app.post("/api/tasks/{tid}/comments")
def create_comment(tid: int, data: dict, u: User = Depends(require_user), s: Session = Depends(get_session)):
    t = s.get(Task, tid)
    if not t: raise HTTPException(404)
    text = data.get("text", "").strip()
    if not text: raise HTTPException(400, "Texte requis")
    c = Comment(task_id=tid, user_id=u.id, text=text)
    s.add(c)
    # Notifications pour les personnes mentionnées (@)
    mentions = data.get("mentions") or []
    notified = set()
    for mid in mentions:
        try: mid = int(mid)
        except (TypeError, ValueError): continue
        if mid in notified or mid == u.id: continue
        if s.get(User, mid):
            _notify(s, mid, "mention",
                    f"{u.name} t'a mentionné",
                    f"Sur la tâche « {t.title} » :\n{text[:200]}",
                    task_id=tid, actor_id=u.id)
            notified.add(mid)
    s.commit(); s.refresh(c)
    return {"id": c.id, "author": u.name, "text": c.text, "created_at": c.created_at.isoformat()}

@app.delete("/api/comments/{cid}")
def delete_comment(cid: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    c = s.get(Comment, cid)
    if c:
        if u.role != "admin" and c.user_id != u.id: raise HTTPException(403)
        s.delete(c); s.commit()
    return {"ok": True}

# ---------------- API : Recherche ----------------
@app.get("/api/search")
def search_api(q: str = "", u: User = Depends(require_user), s: Session = Depends(get_session)):
    if not q.strip() or len(q.strip()) < 2: return {"tasks": [], "projects": []}
    ql = q.strip().lower()
    tasks = [task_dict(t) for t in s.exec(select(Task)).all()
             if ql in t.title.lower() or ql in (t.description or "").lower()][:15]
    projects = [{"id": p.id, "name": p.name} for p in s.exec(select(Project)).all()
                if ql in p.name.lower()][:5]
    return {"tasks": tasks, "projects": projects}

# ---------------- API : Tags ----------------
@app.get("/api/projects/{pid}/tags")
def list_tags(pid: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    return [{"id": t.id, "name": t.name, "color": t.color}
            for t in s.exec(select(Tag).where(Tag.project_id == pid)).all()]

@app.post("/api/projects/{pid}/tags")
def create_tag(pid: int, data: dict, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    name = data.get("name", "").strip()
    if not name: raise HTTPException(400, "Nom requis")
    t = Tag(project_id=pid, name=name, color=data.get("color", "#e8642f"))
    s.add(t); s.commit(); s.refresh(t)
    return {"id": t.id, "name": t.name, "color": t.color}

@app.delete("/api/tags/{tag_id}")
def delete_tag(tag_id: int, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    t = s.get(Tag, tag_id)
    if t:
        for tt in s.exec(select(TaskTag).where(TaskTag.tag_id == tag_id)).all(): s.delete(tt)
        s.delete(t); s.commit()
    return {"ok": True}

@app.get("/api/tasks/{tid}/tags")
def get_task_tags(tid: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    tts = s.exec(select(TaskTag).where(TaskTag.task_id == tid)).all()
    tags = [s.get(Tag, tt.tag_id) for tt in tts]
    return [{"id": t.id, "name": t.name, "color": t.color} for t in tags if t]

@app.post("/api/tasks/{tid}/tags/{tag_id}")
def add_task_tag(tid: int, tag_id: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    t = s.get(Task, tid)
    if not t or not _can_edit_task(u, t): raise HTTPException(403)
    if not s.exec(select(TaskTag).where(TaskTag.task_id == tid, TaskTag.tag_id == tag_id)).first():
        s.add(TaskTag(task_id=tid, tag_id=tag_id)); s.commit()
    return {"ok": True}

@app.delete("/api/tasks/{tid}/tags/{tag_id}")
def remove_task_tag(tid: int, tag_id: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    tt = s.exec(select(TaskTag).where(TaskTag.task_id == tid, TaskTag.tag_id == tag_id)).first()
    if tt: s.delete(tt); s.commit()
    return {"ok": True}

# ---------------- API : Jalons ----------------
@app.get("/api/projects/{pid}/milestones")
def list_milestones(pid: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    return [{"id": m.id, "name": m.name, "due_date": m.due_date.isoformat() if m.due_date else None, "description": m.description}
            for m in s.exec(select(Milestone).where(Milestone.project_id == pid).order_by(Milestone.due_date)).all()]

@app.post("/api/projects/{pid}/milestones")
def create_milestone(pid: int, data: dict, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    name = data.get("name", "").strip()
    if not name: raise HTTPException(400, "Nom requis")
    m = Milestone(project_id=pid, name=name, due_date=_parse_date(data.get("due_date")), description=data.get("description",""))
    s.add(m); s.commit(); s.refresh(m)
    return {"id": m.id, "name": m.name, "due_date": m.due_date.isoformat() if m.due_date else None}

@app.put("/api/milestones/{mid}")
def update_milestone(mid: int, data: dict, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    m = s.get(Milestone, mid)
    if not m: raise HTTPException(404)
    if "name" in data: m.name = data["name"].strip()
    if "due_date" in data: m.due_date = _parse_date(data["due_date"])
    if "description" in data: m.description = data["description"]
    s.add(m); s.commit()
    return {"ok": True}

@app.delete("/api/milestones/{mid}")
def delete_milestone(mid: int, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    m = s.get(Milestone, mid)
    if m:
        for t in s.exec(select(Task).where(Task.milestone_id == mid)).all():
            t.milestone_id = None; s.add(t)
        s.delete(m); s.commit()
    return {"ok": True}

# ---------------- API : Templates (E1) ----------------
@app.get("/api/projects/{pid}/template")
def export_template(pid: int, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    p = s.get(Project, pid)
    if not p: raise HTTPException(404)
    tasks = [{"title": t.title, "description": t.description, "priority": t.priority,
              "status": "todo", "progress": 0} for t in s.exec(select(Task).where(Task.project_id == pid)).all()]
    return {"name": p.name, "description": p.description, "tasks": tasks}

@app.post("/api/projects/from-template")
def create_from_template(data: dict, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    name = data.get("name", "").strip()
    if not name: raise HTTPException(400, "Nom requis")
    p = Project(name=name, description=data.get("description", ""))
    s.add(p); s.commit(); s.refresh(p)
    for td in data.get("tasks", []):
        t = Task(project_id=p.id, title=td.get("title",""), description=td.get("description",""),
                 priority=td.get("priority","m"), status="todo", progress=0)
        s.add(t)
    s.commit()
    return {"id": p.id, "name": p.name}

# ---------------- API : Absences ----------------
@app.get("/api/absences")
def list_absences(u: User = Depends(require_user), s: Session = Depends(get_session)):
    return [{"id": a.id, "user_id": a.user_id, "kind": a.kind,
             "from_date": a.from_date.isoformat(), "to_date": a.to_date.isoformat()}
            for a in s.exec(select(Absence).order_by(Absence.from_date.desc())).all()]

@app.post("/api/absences")
def create_absence(data: dict, u: User = Depends(require_user), s: Session = Depends(get_session)):
    uid = data.get("user_id") or u.id
    # Non-admin : ne peut déclarer que pour soi-même
    if u.role != "admin": uid = u.id
    a = Absence(user_id=uid, kind=data.get("kind", "Congé"),
                from_date=_parse_date(data["from_date"]),
                to_date=_parse_date(data["to_date"]))
    if a.to_date < a.from_date:
        raise HTTPException(400, "La date de fin doit suivre la date de début")
    s.add(a); s.commit(); s.refresh(a)
    return {"id": a.id}

@app.delete("/api/absences/{aid}")
def delete_absence(aid: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    a = s.get(Absence, aid)
    if a:
        if u.role != "admin" and a.user_id != u.id:
            raise HTTPException(403)
        s.delete(a); s.commit()
    return {"ok": True}

# ---------------- API : Alerts ----------------
@app.get("/api/alerts")
def get_alerts(project_id: Optional[int] = None,
               u: User = Depends(require_user), s: Session = Depends(get_session)):
    today = date.today()
    q = select(Task)
    if project_id is not None:
        q = q.where(Task.project_id == project_id)
    tasks = s.exec(q).all()
    users = {x.id: x for x in s.exec(select(User)).all()}
    absences = s.exec(select(Absence)).all()
    acked_keys = {a.alert_key for a in s.exec(select(AckedAlert)).all()}
    def is_absent_now(uid):
        return any(a.user_id == uid and a.from_date <= today <= a.to_date for a in absences)
    out = []
    for t in tasks:
        assignee = users.get(t.assignee_id) if t.assignee_id else None
        aname = assignee.name if assignee else "Non assigné"
        if t.status != "done" and t.due_date and t.due_date < today:
            d = (today - t.due_date).days
            key = f"late:{t.id}"
            if key not in acked_keys:
                out.append({"key": key, "kind": "late", "task_id": t.id, "type": "bad",
                            "title": f"Retard : {t.title}",
                            "msg": f"En retard de {d} jour(s). Assigné à {aname}.",
                            "assignee_email": assignee.email if assignee else None})
        elif t.status != "done" and t.due_date:
            d = (t.due_date - today).days
            if 0 <= d <= 2:
                key = f"soon:{t.id}"
                if key not in acked_keys:
                    out.append({"key": key, "kind": "soon", "task_id": t.id, "type": "warn",
                                "title": f"Échéance proche : {t.title}",
                                "msg": f"À rendre dans {d} jour(s) — {aname}.",
                                "assignee_email": assignee.email if assignee else None})
        if t.status != "done" and t.assignee_id and is_absent_now(t.assignee_id):
            key = f"absent:{t.id}"
            if key not in acked_keys:
                out.append({"key": key, "kind": "absent", "task_id": t.id, "type": "info",
                            "title": f"Personne absente : {t.title}",
                            "msg": f"{aname} est absent(e) aujourd'hui mais a une tâche active.",
                            "assignee_email": assignee.email if assignee else None})
    return out

@app.post("/api/alerts/ack")
def ack_alert(data: dict, u: User = Depends(require_user), s: Session = Depends(get_session)):
    key = data.get("key")
    if not key: raise HTTPException(400)
    if not s.exec(select(AckedAlert).where(AckedAlert.alert_key == key)).first():
        s.add(AckedAlert(alert_key=key)); s.commit()
    return {"ok": True}

@app.post("/api/alerts/ack_all")
def ack_all(u: User = Depends(require_user), s: Session = Depends(get_session)):
    alerts = get_alerts(None, u, s)
    existing = {a.alert_key for a in s.exec(select(AckedAlert)).all()}
    for a in alerts:
        if a["key"] not in existing:
            s.add(AckedAlert(alert_key=a["key"]))
    s.commit()
    return {"ok": True, "count": len(alerts)}

# ---------------- API : Notifications in-app ----------------
@app.get("/api/notifications")
def list_notifications(u: User = Depends(require_user), s: Session = Depends(get_session)):
    rows = s.exec(select(Notification).where(Notification.user_id == u.id)
                 .order_by(Notification.created_at.desc()).limit(40)).all()
    unread = sum(1 for n in rows if not n.read)
    return {"unread": unread, "items": [{
        "id": n.id, "kind": n.kind, "title": n.title, "body": n.body,
        "doc_id": n.doc_id, "task_id": n.task_id, "read": n.read,
        "created_at": n.created_at.isoformat(),
    } for n in rows]}

@app.post("/api/notifications/read")
def mark_notifications_read(data: dict = None, u: User = Depends(require_user), s: Session = Depends(get_session)):
    data = data or {}
    nid = data.get("id")
    if nid:
        s.execute(sa_update(Notification).where(Notification.id == nid, Notification.user_id == u.id).values(read=True))
    else:
        s.execute(sa_update(Notification).where(Notification.user_id == u.id, Notification.read == False).values(read=True))
    s.commit()
    return {"ok": True}

# ---------------- API : Settings (branding) ----------------
@app.get("/api/settings")
def get_settings(s: Session = Depends(get_session)):
    return {"app_name": _setting(s, "app_name", "Helix"),
            "app_logo": _setting(s, "app_logo", "")}

@app.put("/api/settings")
def update_settings(data: dict, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    if "app_name" in data: _set_setting(s, "app_name", data["app_name"].strip() or "Helix")
    if "app_logo" in data: _set_setting(s, "app_logo", data["app_logo"])
    return {"ok": True}

# ---------------- API : Audit ----------------
@app.get("/api/audit")
def get_audit(limit: int = 200, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    logs = s.exec(select(AuditLog).order_by(AuditLog.created_at.desc()).limit(limit)).all()
    return [{"id": l.id, "user_name": l.user_name, "action": l.action,
             "details": l.details, "created_at": l.created_at.isoformat()} for l in logs]

# ---------------- Email ----------------
# Fournisseur 1 (recommandé) : Resend (API HTTP, fonctionne sur Render)
RESEND_API_KEY = os.environ.get("RESEND_API_KEY", "")
EMAIL_FROM     = os.environ.get("EMAIL_FROM", "Helix <onboarding@resend.dev>")
# Fournisseur 2 (repli) : SMTP classique
SMTP_FROM    = os.environ.get("SMTP_FROM",    "manufacturingengineeringteam@alivedx.com")
SMTP_SERVER  = os.environ.get("SMTP_SERVER",  "smtp.office365.com")
SMTP_PORT    = int(os.environ.get("SMTP_PORT", "587"))
SMTP_USER    = os.environ.get("SMTP_USER",    SMTP_FROM)
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD", "")

def email_configured() -> bool:
    return bool(RESEND_API_KEY or SMTP_PASSWORD)

def _send_email(to: str, subject: str, body: str):
    # 1) Resend (HTTP) — prioritaire
    if RESEND_API_KEY:
        import urllib.request
        payload = json.dumps({
            "from": EMAIL_FROM, "to": [to], "subject": subject, "text": body,
        }).encode("utf-8")
        req = urllib.request.Request(
            "https://api.resend.com/emails", data=payload, method="POST",
            headers={"Authorization": f"Bearer {RESEND_API_KEY}", "Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=15) as r:
            r.read()
        return
    # 2) SMTP — repli
    if SMTP_PASSWORD:
        msg = MIMEMultipart("alternative")
        msg["From"] = SMTP_FROM; msg["To"] = to; msg["Subject"] = subject
        msg.attach(MIMEText(body, "plain", "utf-8"))
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT, timeout=10) as s:
            s.ehlo(); s.starttls(); s.ehlo()
            s.login(SMTP_USER, SMTP_PASSWORD)
            s.sendmail(SMTP_FROM, [to], msg.as_string())
        return
    raise ValueError("Aucun fournisseur email configuré (RESEND_API_KEY ou SMTP_PASSWORD).")

@app.post("/api/send-email")
def send_email_endpoint(data: dict, u: User = Depends(require_user)):
    to      = data.get("to", "").strip()
    subject = data.get("subject", "").strip()
    body    = data.get("body", "").strip()
    if not to or not subject or not body:
        raise HTTPException(400, "Champs manquants (to, subject, body)")
    try:
        _send_email(to, subject, body)
    except ValueError as e:
        raise HTTPException(503, str(e))
    except Exception as e:
        raise HTTPException(502, f"Échec envoi email : {e}")
    return {"ok": True}

# ---------------- API : IA — estimation intelligente de charge ----------------
AI_MODEL = os.environ.get("AI_MODEL", "claude-opus-4-8")  # configurable (ex: claude-sonnet-4-6)
_ai_client = None

def _get_ai_client():
    global _ai_client
    if anthropic is None:
        raise HTTPException(503, "Le SDK Anthropic n'est pas installé (redéploie avec le requirements à jour).")
    if not os.environ.get("ANTHROPIC_API_KEY"):
        raise HTTPException(503, "IA non configurée : ajoute la variable ANTHROPIC_API_KEY sur Render.")
    if _ai_client is None:
        _ai_client = anthropic.Anthropic()  # lit ANTHROPIC_API_KEY depuis l'environnement
    return _ai_client

_AI_SYSTEM = (
    "Tu es un chef de projet industriel expert en estimation de charge de travail. "
    "On te donne une liste de tâches (titre, description, priorité). Pour CHAQUE tâche, "
    "estime le nombre d'heures de travail réalistes nécessaires pour la réaliser entièrement.\n"
    "Repères : tâche très simple 1-2h, simple 3-4h, moyenne 6-12h, complexe 16-24h, "
    "très complexe 32-60h. Tiens compte de la description et de la priorité (une priorité haute "
    "n'augmente pas forcément la charge, mais l'urgence si). Sois pragmatique et cohérent entre les tâches. "
    "Donne un 'rationale' très court (8 mots max). Réponds uniquement via le format structuré demandé."
)

_AI_SCHEMA = {
    "type": "object",
    "properties": {
        "estimates": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "id": {"type": "integer"},
                    "hours": {"type": "number"},
                    "rationale": {"type": "string"},
                },
                "required": ["id", "hours", "rationale"],
                "additionalProperties": False,
            },
        }
    },
    "required": ["estimates"],
    "additionalProperties": False,
}

@app.post("/api/ai/estimate-load")
def ai_estimate_load(data: dict, u: User = Depends(require_admin), s: Session = Depends(get_session)):
    project_id = data.get("project_id")
    q = select(Task).where(Task.status != "done")
    if project_id:
        q = q.where(Task.project_id == project_id)
    tasks = s.exec(q).all()
    to_estimate = [t for t in tasks if not t.estimated_hours or t.estimated_hours <= 0][:100]
    if not to_estimate:
        return {"ok": True, "updated": 0, "message": "Toutes les tâches actives ont déjà une estimation."}

    client = _get_ai_client()
    prio_lbl = {"h": "haute", "m": "moyenne", "l": "basse"}
    payload = [{
        "id": t.id, "titre": t.title,
        "description": (t.description or "")[:400],
        "priorite": prio_lbl.get(t.priority, t.priority),
    } for t in to_estimate]
    user_msg = ("Estime la charge en heures pour ces tâches :\n"
                + json.dumps(payload, ensure_ascii=False, indent=1))

    try:
        resp = client.messages.create(
            model=AI_MODEL,
            max_tokens=8000,
            system=[{"type": "text", "text": _AI_SYSTEM, "cache_control": {"type": "ephemeral"}}],
            messages=[{"role": "user", "content": user_msg}],
            output_config={"format": {"type": "json_schema", "schema": _AI_SCHEMA}},
        )
    except Exception as e:
        msg = str(e)
        if "authentication" in msg.lower() or "x-api-key" in msg.lower() or "401" in msg:
            raise HTTPException(401, "Clé API Anthropic invalide (vérifie ANTHROPIC_API_KEY).")
        raise HTTPException(502, f"Erreur lors de l'analyse IA : {msg}")

    text = next((b.text for b in resp.content if getattr(b, "type", None) == "text"), "")
    try:
        estimates = json.loads(text).get("estimates", [])
    except Exception:
        raise HTTPException(502, "Réponse IA illisible.")

    valid_ids = {t.id for t in to_estimate}
    updated, applied = 0, []
    for e in estimates:
        tid = e.get("id")
        hrs = e.get("hours")
        if tid in valid_ids and isinstance(hrs, (int, float)) and hrs > 0:
            hrs = max(0.5, round(float(hrs) * 2) / 2)  # arrondi au 0.5h, plancher 0.5
            s.execute(sa_update(Task).where(Task.id == tid).values(estimated_hours=hrs))
            updated += 1
            applied.append({"id": tid, "hours": hrs, "rationale": e.get("rationale", "")})
    _audit(s, u.name, "Estimation IA", f"{updated} tâche(s) estimée(s) (modèle {AI_MODEL})")
    s.commit()
    return {"ok": True, "updated": updated, "applied": applied,
            "model": AI_MODEL,
            "usage": {"input": resp.usage.input_tokens, "output": resp.usage.output_tokens}}

# ---------------- API : Gestion documentaire (qualité) ----------------
DOC_MAX_SIZE = 25 * 1024 * 1024  # 25 Mo
DOC_ALLOWED_EXT = (".docx", ".doc", ".pdf", ".xlsx", ".xls", ".pptx", ".ppt", ".txt", ".md", ".csv")
DOC_STATUS = ("draft", "review", "approved")
# Workflow documentaire (style Veeva) — phases ordonnées avant la validation officielle QMS
DOC_PHASES = ["redaction", "revue_equipe", "revue_qa", "approbation", "pret_qms"]
DOC_PHASE_STATUS = {  # statut "simple" dérivé de la phase (rétrocompat)
    "redaction": "draft", "revue_equipe": "review", "revue_qa": "review",
    "approbation": "review", "pret_qms": "approved",
}
DOC_PHASE_LABELS = {
    "redaction": "Rédaction", "revue_equipe": "Revue équipe", "revue_qa": "Revue QA",
    "approbation": "Approbation", "pret_qms": "Prêt pour QMS",
}
# Types de documents (préfixe de référence)
DOC_TYPES = {
    "SOP": "SOP", "PROTO": "Protocole", "REPORT": "Rapport",
    "FORM": "Formulaire", "IT": "Instruction", "DOC": "Document",
}
# Délai indicatif max par phase (jours) — au-delà : alerte SLA (goulot)
DOC_SLA_DAYS = {
    "redaction": 14, "revue_equipe": 5, "revue_qa": 7, "approbation": 5, "pret_qms": 0,
}
# Phases nécessitant une signature électronique (mot de passe + motif)
DOC_SIGN_PHASES = ("approbation", "pret_qms")

def _next_doc_reference(s: Session, doc_type: str) -> str:
    prefix = doc_type if doc_type in DOC_TYPES else "DOC"
    year = datetime.utcnow().year
    pat = f"{prefix}-{year}-"
    existing = s.exec(select(Document).where(Document.reference.startswith(pat))).all()
    nums = []
    for d in existing:
        try: nums.append(int((d.reference or "").rsplit("-", 1)[-1]))
        except (ValueError, IndexError): pass
    nxt = (max(nums) + 1) if nums else 1
    return f"{pat}{nxt:03d}"

def _user_name(s: Session, uid):
    if not uid: return None
    u = s.get(User, uid)
    return u.name if u else None

def _doc_dict(s: Session, d: Document, versions=None, me_id=None):
    if versions is None:
        versions = s.exec(select(DocumentVersion).where(DocumentVersion.document_id == d.id)
                          .order_by(DocumentVersion.version.desc())).all()
    cur = versions[0] if versions else None
    cur_ver = cur.version if cur else 0
    phase = d.phase or "redaction"
    # SLA : jours dans la phase actuelle
    since = d.phase_since or d.created_at
    days_in_phase = (datetime.utcnow() - since).days if since else 0
    sla_days = DOC_SLA_DAYS.get(phase, 0)
    sla_over = bool(sla_days and days_in_phase > sla_days)
    # Signatures + accusés de lecture (version courante)
    sign_count = len(s.exec(select(DocSignature).where(DocSignature.document_id == d.id)).all())
    acks = s.exec(select(DocAck).where(DocAck.document_id == d.id, DocAck.version == cur_ver)).all()
    my_ack = bool(me_id and any(a.user_id == me_id for a in acks))
    return {
        "id": d.id, "name": d.name, "description": d.description,
        "doc_type": d.doc_type or "DOC", "reference": d.reference or "",
        "status": d.status, "project_id": d.project_id,
        "phase": phase,
        "phase_since": since.isoformat() if since else None,
        "days_in_phase": days_in_phase, "sla_days": sla_days, "sla_over": sla_over,
        "assigned_to": d.assigned_to, "assigned_to_name": _user_name(s, d.assigned_to),
        "locked_by": d.locked_by, "locked_by_name": _user_name(s, d.locked_by),
        "locked_at": d.locked_at.isoformat() if d.locked_at else None,
        "created_by": d.created_by, "created_by_name": _user_name(s, d.created_by),
        "created_at": d.created_at.isoformat() if d.created_at else None,
        "version_count": len(versions),
        "last_version": cur_ver,
        "last_modified_by": _user_name(s, cur.uploaded_by) if cur else None,
        "last_modified_at": cur.uploaded_at.isoformat() if cur and cur.uploaded_at else None,
        "last_filename": cur.filename if cur else None,
        "sign_count": sign_count,
        "ack_count": len(acks), "my_ack": my_ack,
        "needs_ack": phase == "pret_qms",
    }

@app.get("/api/documents")
def list_documents(u: User = Depends(require_user), s: Session = Depends(get_session)):
    docs = s.exec(select(Document).order_by(Document.name)).all()
    return [_doc_dict(s, d, me_id=u.id) for d in docs]

@app.get("/api/documents/{doc_id}")
def get_document(doc_id: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    d = s.get(Document, doc_id)
    if not d: raise HTTPException(404, "Document introuvable")
    versions = s.exec(select(DocumentVersion).where(DocumentVersion.document_id == doc_id)
                     .order_by(DocumentVersion.version.desc())).all()
    out = _doc_dict(s, d, versions, me_id=u.id)
    out["versions"] = [{
        "id": v.id, "version": v.version, "filename": v.filename,
        "size": v.size, "note": v.note,
        "uploaded_by_name": _user_name(s, v.uploaded_by),
        "uploaded_at": v.uploaded_at.isoformat() if v.uploaded_at else None,
    } for v in versions]
    wf = s.exec(select(DocWorkflowEvent).where(DocWorkflowEvent.document_id == doc_id)
               .order_by(DocWorkflowEvent.created_at.asc())).all()
    out["workflow"] = [{
        "phase": w.phase, "assigned_to_name": _user_name(s, w.assigned_to),
        "moved_by_name": _user_name(s, w.moved_by), "note": w.note,
        "created_at": w.created_at.isoformat() if w.created_at else None,
    } for w in wf]
    sigs = s.exec(select(DocSignature).where(DocSignature.document_id == doc_id)
                 .order_by(DocSignature.signed_at.asc())).all()
    out["signatures"] = [{
        "user_name": sg.user_name, "meaning": sg.meaning, "reason": sg.reason,
        "version": sg.version, "signed_at": sg.signed_at.isoformat() if sg.signed_at else None,
    } for sg in sigs]
    cur_ver = out["last_version"]
    acks = s.exec(select(DocAck).where(DocAck.document_id == doc_id, DocAck.version == cur_ver)).all()
    out["acks"] = [{
        "user_name": _user_name(s, a.user_id),
        "acknowledged_at": a.acknowledged_at.isoformat() if a.acknowledged_at else None,
    } for a in acks]
    return out

@app.post("/api/documents/{doc_id}/transition")
def transition_document(doc_id: int, data: dict, u: User = Depends(require_user), s: Session = Depends(get_session)):
    d = s.get(Document, doc_id)
    if not d: raise HTTPException(404, "Document introuvable")
    phase = data.get("phase")
    if phase not in DOC_PHASES:
        raise HTTPException(400, "Phase invalide")
    assigned_to = data.get("assigned_to")
    if assigned_to:
        if not s.get(User, assigned_to):
            raise HTTPException(400, "Personne assignée introuvable")
    else:
        assigned_to = None
    note = (data.get("note") or "").strip()
    phase_lbl = DOC_PHASE_LABELS.get(phase, phase)
    # --- Signature électronique obligatoire (esprit 21 CFR Part 11) ---
    if phase in DOC_SIGN_PHASES:
        password = data.get("password") or ""
        reason = (data.get("reason") or "").strip()
        if not password or not reason:
            raise HTTPException(400, "Signature requise : mot de passe et motif obligatoires pour cette phase.")
        if not bcrypt.verify(password, u.password_hash):
            raise HTTPException(403, "Signature refusée : mot de passe incorrect.")
        cur = s.exec(select(DocumentVersion).where(DocumentVersion.document_id == doc_id)
                    .order_by(DocumentVersion.version.desc())).first()
        s.add(DocSignature(document_id=doc_id, user_id=u.id, user_name=u.name,
                           phase=phase, meaning=phase_lbl, reason=reason,
                           version=cur.version if cur else 0))
        _audit(s, u.name, "Signature électronique", f"{d.name} — {phase_lbl} : {reason}")
    s.execute(sa_update(Document).where(Document.id == doc_id).values(
        phase=phase, assigned_to=assigned_to, phase_since=datetime.utcnow(),
        status=DOC_PHASE_STATUS.get(phase, d.status)))
    s.add(DocWorkflowEvent(document_id=doc_id, phase=phase, assigned_to=assigned_to,
                           moved_by=u.id, note=note))
    _audit(s, u.name, "Transition document", f"{d.name} → {phase}"
           + (f" (chez {_user_name(s, assigned_to)})" if assigned_to else ""))
    # Notification in-app pour la personne assignée
    if assigned_to and assigned_to != u.id:
        _notify(s, assigned_to, "doc_assigned",
                f"Document assigné : {d.name}",
                f"« {d.name} » est passé en phase « {phase_lbl} » et requiert ton action."
                + (f"\nNote : {note}" if note else ""),
                doc_id=doc_id, actor_id=u.id)
    s.commit()
    # L'email part via Outlook (mailto) côté client ; ici on renvoie juste les infos destinataire.
    target = s.get(User, assigned_to) if assigned_to else None
    return {"ok": True, "phase": phase,
            "assignee_email": target.email if target else None,
            "assignee_name": target.name if target else None}

@app.post("/api/documents/{doc_id}/ack")
def ack_document(doc_id: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    d = s.get(Document, doc_id)
    if not d: raise HTTPException(404, "Document introuvable")
    cur = s.exec(select(DocumentVersion).where(DocumentVersion.document_id == doc_id)
                .order_by(DocumentVersion.version.desc())).first()
    ver = cur.version if cur else 0
    existing = s.exec(select(DocAck).where(DocAck.document_id == doc_id,
                      DocAck.user_id == u.id, DocAck.version == ver)).first()
    if not existing:
        s.add(DocAck(document_id=doc_id, user_id=u.id, version=ver))
        _audit(s, u.name, "Accusé de lecture", f"{d.name} (v{ver})")
        s.commit()
    return {"ok": True}

@app.post("/api/documents")
async def create_document(
    name: str = Form(...), description: str = Form(""),
    project_id: Optional[int] = Form(None), note: str = Form(""),
    doc_type: str = Form("DOC"),
    file: UploadFile = File(...),
    u: User = Depends(require_user), s: Session = Depends(get_session)):
    name = name.strip()
    if not name: raise HTTPException(400, "Nom requis")
    fname = (file.filename or "").lower()
    if not fname.endswith(DOC_ALLOWED_EXT):
        raise HTTPException(400, f"Format non supporté ({', '.join(DOC_ALLOWED_EXT)})")
    content = await file.read()
    if len(content) > DOC_MAX_SIZE:
        raise HTTPException(400, "Fichier trop volumineux (max 25 Mo)")
    dtype = doc_type if doc_type in DOC_TYPES else "DOC"
    reference = _next_doc_reference(s, dtype)
    d = Document(name=name, description=description.strip(),
                 project_id=project_id, created_by=u.id, status="draft",
                 phase="redaction", assigned_to=u.id,
                 doc_type=dtype, reference=reference, phase_since=datetime.utcnow())
    s.add(d); s.commit(); s.refresh(d)
    v = DocumentVersion(document_id=d.id, version=1, filename=file.filename,
                        mime_type=file.content_type or "application/octet-stream",
                        size=len(content), content=content, uploaded_by=u.id,
                        note=note.strip() or "Version initiale")
    s.add(v)
    s.add(DocWorkflowEvent(document_id=d.id, phase="redaction", assigned_to=u.id,
                           moved_by=u.id, note="Création du document"))
    _audit(s, u.name, "Création document", f"{name} (v1)")
    s.commit()
    return _doc_dict(s, d)

@app.post("/api/documents/{doc_id}/versions")
async def upload_version(
    doc_id: int, note: str = Form(""), file: UploadFile = File(...),
    u: User = Depends(require_user), s: Session = Depends(get_session)):
    d = s.get(Document, doc_id)
    if not d: raise HTTPException(404, "Document introuvable")
    # Si verrouillé par quelqu'un d'autre, refuser (sauf admin)
    if d.locked_by and d.locked_by != u.id and u.role != "admin":
        raise HTTPException(403, f"Document verrouillé par {_user_name(s, d.locked_by)}")
    fname = (file.filename or "").lower()
    if not fname.endswith(DOC_ALLOWED_EXT):
        raise HTTPException(400, f"Format non supporté ({', '.join(DOC_ALLOWED_EXT)})")
    content = await file.read()
    if len(content) > DOC_MAX_SIZE:
        raise HTTPException(400, "Fichier trop volumineux (max 25 Mo)")
    last = s.exec(select(DocumentVersion).where(DocumentVersion.document_id == doc_id)
                 .order_by(DocumentVersion.version.desc())).first()
    next_v = (last.version + 1) if last else 1
    v = DocumentVersion(document_id=doc_id, version=next_v, filename=file.filename,
                        mime_type=file.content_type or "application/octet-stream",
                        size=len(content), content=content, uploaded_by=u.id,
                        note=note.strip())
    s.add(v)
    # libère le verrou et enregistre
    s.execute(sa_update(Document).where(Document.id == doc_id).values(locked_by=None, locked_at=None))
    _audit(s, u.name, "Nouvelle version document", f"{d.name} (v{next_v})")
    s.commit()
    return _doc_dict(s, s.get(Document, doc_id))

@app.get("/api/documents/{doc_id}/versions/{ver_id}/download")
def download_version(doc_id: int, ver_id: int,
                     u: User = Depends(require_user), s: Session = Depends(get_session)):
    v = s.get(DocumentVersion, ver_id)
    if not v or v.document_id != doc_id:
        raise HTTPException(404, "Version introuvable")
    headers = {"Content-Disposition": f"attachment; filename*=UTF-8''{_urlquote(v.filename)}"}
    return StreamingResponse(io.BytesIO(v.content),
                             media_type=v.mime_type or "application/octet-stream",
                             headers=headers)

@app.get("/api/documents/{doc_id}/versions/{ver_id}/view")
def view_version(doc_id: int, ver_id: int,
                 u: User = Depends(require_user), s: Session = Depends(get_session)):
    """Sert le fichier en ligne (inline) pour l'aperçu dans le navigateur (PDF, images…)."""
    v = s.get(DocumentVersion, ver_id)
    if not v or v.document_id != doc_id:
        raise HTTPException(404, "Version introuvable")
    headers = {"Content-Disposition": f"inline; filename*=UTF-8''{_urlquote(v.filename)}"}
    return StreamingResponse(io.BytesIO(v.content),
                             media_type=v.mime_type or "application/octet-stream",
                             headers=headers)

# ---------------- API : Commentaires de documents ----------------
@app.get("/api/documents/{doc_id}/comments")
def list_doc_comments(doc_id: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    users = {x.id: x.name for x in s.exec(select(User)).all()}
    rows = s.exec(select(DocComment).where(DocComment.document_id == doc_id)
                 .order_by(DocComment.created_at)).all()
    return [{"id": c.id, "user_id": c.user_id, "author": users.get(c.user_id, "?"),
             "text": c.text, "created_at": c.created_at.isoformat()} for c in rows]

@app.post("/api/documents/{doc_id}/comments")
def create_doc_comment(doc_id: int, data: dict, u: User = Depends(require_user), s: Session = Depends(get_session)):
    d = s.get(Document, doc_id)
    if not d: raise HTTPException(404, "Document introuvable")
    text = (data.get("text") or "").strip()
    if not text: raise HTTPException(400, "Texte requis")
    c = DocComment(document_id=doc_id, user_id=u.id, text=text)
    s.add(c)
    notified = set()
    for mid in (data.get("mentions") or []):
        try: mid = int(mid)
        except (TypeError, ValueError): continue
        if mid in notified or mid == u.id: continue
        if s.get(User, mid):
            _notify(s, mid, "mention", f"{u.name} t'a mentionné",
                    f"Sur le document « {d.name} » :\n{text[:200]}",
                    doc_id=doc_id, actor_id=u.id)
            notified.add(mid)
    s.commit(); s.refresh(c)
    return {"id": c.id, "author": u.name, "text": c.text, "created_at": c.created_at.isoformat()}

@app.delete("/api/doc-comments/{cid}")
def delete_doc_comment(cid: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    c = s.get(DocComment, cid)
    if c:
        if u.role != "admin" and c.user_id != u.id: raise HTTPException(403)
        s.delete(c); s.commit()
    return {"ok": True}

@app.post("/api/documents/{doc_id}/lock")
def lock_document(doc_id: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    d = s.get(Document, doc_id)
    if not d: raise HTTPException(404, "Document introuvable")
    if d.locked_by and d.locked_by != u.id:
        raise HTTPException(409, f"Déjà verrouillé par {_user_name(s, d.locked_by)}")
    s.execute(sa_update(Document).where(Document.id == doc_id)
              .values(locked_by=u.id, locked_at=datetime.utcnow()))
    _audit(s, u.name, "Verrouillage document", d.name)
    s.commit()
    return {"ok": True}

@app.post("/api/documents/{doc_id}/unlock")
def unlock_document(doc_id: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    d = s.get(Document, doc_id)
    if not d: raise HTTPException(404, "Document introuvable")
    if d.locked_by and d.locked_by != u.id and u.role != "admin":
        raise HTTPException(403, "Seul le détenteur du verrou ou un admin peut déverrouiller")
    s.execute(sa_update(Document).where(Document.id == doc_id).values(locked_by=None, locked_at=None))
    _audit(s, u.name, "Déverrouillage document", d.name)
    s.commit()
    return {"ok": True}

@app.put("/api/documents/{doc_id}")
def update_document(doc_id: int, data: dict, u: User = Depends(require_user), s: Session = Depends(get_session)):
    d = s.get(Document, doc_id)
    if not d: raise HTTPException(404, "Document introuvable")
    updates = {}
    if "name" in data and data["name"].strip(): updates["name"] = data["name"].strip()
    if "description" in data: updates["description"] = data["description"].strip()
    if "status" in data and data["status"] in DOC_STATUS: updates["status"] = data["status"]
    if updates:
        s.execute(sa_update(Document).where(Document.id == doc_id).values(**updates))
        _audit(s, u.name, "Modification document", f"{d.name} — {list(updates.keys())}")
        s.commit()
    return {"ok": True}

@app.delete("/api/documents/{doc_id}")
def delete_document(doc_id: int, u: User = Depends(require_user), s: Session = Depends(get_session)):
    d = s.get(Document, doc_id)
    if d:
        if u.role != "admin" and d.created_by != u.id:
            raise HTTPException(403, "Seul le créateur ou un admin peut supprimer ce document")
        for v in s.exec(select(DocumentVersion).where(DocumentVersion.document_id == doc_id)).all():
            s.delete(v)
        for w in s.exec(select(DocWorkflowEvent).where(DocWorkflowEvent.document_id == doc_id)).all():
            s.delete(w)
        for sg in s.exec(select(DocSignature).where(DocSignature.document_id == doc_id)).all():
            s.delete(sg)
        for a in s.exec(select(DocAck).where(DocAck.document_id == doc_id)).all():
            s.delete(a)
        for cm in s.exec(select(DocComment).where(DocComment.document_id == doc_id)).all():
            s.delete(cm)
        _audit(s, u.name, "Suppression document", d.name)
        s.delete(d); s.commit()
    return {"ok": True}

# ---------------- API : Document Import ----------------
@app.post("/api/parse-document")
async def parse_document(file: UploadFile = File(...), u: User = Depends(require_user)):
    """Extrait le texte d'un PDF/DOCX/TXT et le renvoie pour analyse côté navigateur."""
    name = file.filename.lower()
    content = await file.read()
    text = ""
    try:
        if name.endswith(".txt") or name.endswith(".md"):
            text = content.decode("utf-8", errors="ignore")
        elif name.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(content))
            text = "\n".join(p.text for p in doc.paragraphs)
        elif name.endswith(".pdf"):
            from pypdf import PdfReader
            r = PdfReader(io.BytesIO(content))
            text = "\n".join(p.extract_text() or "" for p in r.pages)
        else:
            raise HTTPException(400, "Format non supporté (PDF, DOCX, TXT seulement)")
    except Exception as e:
        raise HTTPException(400, f"Lecture impossible : {e}")
    return {"text": text}

# ---------------- Rapport hebdomadaire ----------------
REPORT_RECIPIENT = os.environ.get("WEEKLY_REPORT_EMAIL", "charlotte.foujols@alivedx.com")

def _generate_weekly_report():
    try:
        with Session(engine) as s:
            projects  = s.exec(select(Project)).all()
            all_tasks = s.exec(select(Task)).all()
            users     = {x.id: x.name for x in s.exec(select(User)).all()}
            app_name  = _setting(s, "app_name", "Helix")

        today_str = datetime.utcnow().strftime("%d/%m/%Y")
        body  = f"Rapport hebdomadaire — {today_str}\n"
        body += "=" * 55 + "\n\n"

        total_done = total_all = 0
        for proj in projects:
            tasks = [t for t in all_tasks if t.project_id == proj.id]
            if not tasks:
                continue
            done  = [t for t in tasks if t.status == "done"]
            prog  = [t for t in tasks if t.status == "prog"]
            todo  = [t for t in tasks if t.status == "todo"]
            late  = [t for t in tasks if t.status != "done" and t.due_date and
                     t.due_date.isoformat() < datetime.utcnow().date().isoformat()]
            pct   = round(len(done) / len(tasks) * 100) if tasks else 0
            total_done += len(done); total_all += len(tasks)

            body += f"📁  {proj.name}\n"
            body += f"    Avancement : {pct}%  ({len(done)}/{len(tasks)} terminées)\n"
            body += f"    En cours : {len(prog)}  |  À faire : {len(todo)}  |  Terminées : {len(done)}\n"
            if late:
                body += f"    ⚠  EN RETARD ({len(late)}) :\n"
                for t in late[:10]:
                    body += f"       • {t.title}  —  {users.get(t.assignee_id,'Non assigné')}  (échéance : {t.due_date})\n"
            body += "\n"

        global_pct = round(total_done / total_all * 100) if total_all else 0
        body += "-" * 55 + "\n"
        body += f"GLOBAL : {global_pct}% terminé  ({total_done}/{total_all} tâches)\n\n"
        body += "Ce rapport est généré automatiquement par l'application.\n"

        subject = f"[{app_name}] Rapport hebdomadaire — {today_str}"
        _send_email(REPORT_RECIPIENT, subject, body)
        print(f"[Atelier] Rapport hebdomadaire envoyé à {REPORT_RECIPIENT}")
    except Exception as e:
        print(f"[Atelier] Erreur rapport hebdomadaire : {e}")

# Scheduler — lundi 8h heure de Paris
_scheduler = BackgroundScheduler(timezone="Europe/Paris")
_scheduler.add_job(_generate_weekly_report, "cron", day_of_week="mon", hour=8, minute=0, id="weekly_report")

@app.on_event("startup")
def _start_scheduler():
    _scheduler.start()

@app.on_event("shutdown")
def _stop_scheduler():
    _scheduler.shutdown(wait=False)
